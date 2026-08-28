import os
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def make_doc():
    doc = Document()
    
    # Page setup - Portrait with 0.6 inch margins for maximum table width
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    # Styles
    # Primary Accent: Forest Green (#1B4D3E), Secondary: Emerald (#2E7D32), Dark: #222222
    PRIMARY_COLOR = RGBColor(27, 77, 62)
    SECONDARY_COLOR = RGBColor(46, 125, 50)
    DARK_TEXT = RGBColor(34, 34, 34)
    MUTED_TEXT = RGBColor(100, 100, 100)

    # Document Header
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    p_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run("ECOHARVEST PLATFORM: USER ACCEPTANCE TESTING (UAT) MASTER MANUAL")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = PRIMARY_COLOR

    # Subtitle / Metadata Box
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(12)
    
    runs_meta = [
        ("Document Reference: ", True), ("ECO-UAT-MAN-2026-V1.0  |  ", False),
        ("Classification: ", True), ("Production-Grade Academic QA Standard\n", False),
        ("Lead QA Architect & Principal Systems Auditor: ", True), ("Antigravity Verification Office\n", False),
        ("Target Client Ecosystem: ", True), ("EcoHarvest Mobile App (iOS/Android Expo SDK 57) & Governance Command Center (Expo Web)\n", False),
        ("Backend Services: ", True), ("Node.js/Express (Port 5000), Python Flask AI (Port 5001/5002), Google Gemini 2.5 Flash, Stripe Escrow, Uber Direct Simulator, MongoDB Atlas.", False)
    ]
    for text, bold in runs_meta:
        r = p_sub.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.bold = bold
        r.font.color.rgb = DARK_TEXT if bold else MUTED_TEXT

    # Divider line
    p_div = doc.add_paragraph()
    p_div.paragraph_format.space_after = Pt(8)
    r_div = p_div.add_run("―" * 70)
    r_div.font.color.rgb = RGBColor(200, 200, 200)

    # Executive Summary
    p_exec = doc.add_paragraph()
    p_exec.paragraph_format.space_before = Pt(4)
    p_exec.paragraph_format.space_after = Pt(4)
    r_exec_head = p_exec.add_run("Executive Summary & Test Execution Protocol")
    r_exec_head.font.name = "Arial"
    r_exec_head.font.size = Pt(12)
    r_exec_head.font.bold = True
    r_exec_head.font.color.rgb = SECONDARY_COLOR

    p_exec_body = doc.add_paragraph()
    p_exec_body.paragraph_format.space_after = Pt(8)
    r_eb = p_exec_body.add_run(
        "This User Acceptance Testing (UAT) Case Manual provides the definitive verification blueprint for certifying the EcoHarvest agri-tech platform. "
        "Test cases are organized across distinct functional suites for each user persona, AI grading sub-engine, Stripe escrow pipeline, Uber Direct logistics simulator, and administrator governance tabs.\n"
        "Evaluation Scale: "
    )
    r_eb.font.name = "Arial"
    r_eb.font.size = Pt(9.5)
    
    r_p = p_exec_body.add_run("Pass (P) ")
    r_p.font.bold = True
    r_p.font.color.rgb = SECONDARY_COLOR
    r_p_desc = p_exec_body.add_run("= Expected technical outcome achieved with 0 defects | ")
    r_p_desc.font.size = Pt(9.5)
    
    r_f = p_exec_body.add_run("Fail (F) ")
    r_f.font.bold = True
    r_f.font.color.rgb = RGBColor(198, 40, 40)
    r_f_desc = p_exec_body.add_run("= Logic error, UI crash, API failure, or security defect | ")
    r_f_desc.font.size = Pt(9.5)
    
    r_b = p_exec_body.add_run("Blocked (B) ")
    r_b.font.bold = True
    r_b.font.color.rgb = RGBColor(230, 81, 0)
    r_b_desc = p_exec_body.add_run("= Upstream dependency prevents test execution.")
    r_b_desc.font.size = Pt(9.5)

    # Read markdown to extract suites and tables
    with open("/Users/mevinimunaweera/Documents/EcoHarvest/ECOHARVEST_UAT_TEST_MANUAL.md", "r", encoding="utf-8") as f:
        md_content = f.read()

    # Split into sections
    suite_blocks = re.split(r'\n## (Suite \d+: [^\n]+)\n', md_content)
    
    # Widths in inches (Total ~7.3 inches)
    COL_WIDTHS = [Inches(1.05), Inches(1.35), Inches(2.25), Inches(2.00), Inches(0.65)]
    
    for i in range(1, len(suite_blocks), 2):
        suite_title = suite_blocks[i]
        suite_body = suite_blocks[i+1]
        
        # Add Suite Title
        p_st = doc.add_paragraph()
        p_st.paragraph_format.space_before = Pt(14)
        p_st.paragraph_format.space_after = Pt(6)
        p_st.paragraph_format.keep_with_next = True
        r_st = p_st.add_run(suite_title)
        r_st.font.name = "Arial"
        r_st.font.size = Pt(12.5)
        r_st.font.bold = True
        r_st.font.color.rgb = PRIMARY_COLOR
        
        # Subsections (### 1.1 ...)
        sub_blocks = re.split(r'\n### (\d+\.\d+ [^\n]+)\n', suite_body)
        start_sub = 1 if len(sub_blocks) > 1 else 0
        
        for j in range(1, len(sub_blocks), 2):
            sub_title = sub_blocks[j]
            sub_body = sub_blocks[j+1]
            
            p_subt = doc.add_paragraph()
            p_subt.paragraph_format.space_before = Pt(8)
            p_subt.paragraph_format.space_after = Pt(4)
            p_subt.paragraph_format.keep_with_next = True
            r_subt = p_subt.add_run(sub_title)
            r_subt.font.name = "Arial"
            r_subt.font.size = Pt(10.5)
            r_subt.font.bold = True
            r_subt.font.color.rgb = SECONDARY_COLOR
            
            # Find Markdown Table
            table_match = re.search(r'(\| Test ID \|.*?\n\| :---.*?\n)((?:\| \*\*TC-.*?\n)+)', sub_body)
            if table_match:
                rows_text = table_match.group(2).strip().split('\n')
                
                # Create Word Table
                table = doc.add_table(rows=len(rows_text) + 1, cols=5)
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                set_table_borders(table, color="D0D7DE")
                
                # Header Row
                hdr_cells = table.rows[0].cells
                headers = ["Test ID", "Feature Component", "Action / Instructions", "Expected Technical Outcome", "Pass / Fail"]
                for c_idx, h_text in enumerate(headers):
                    cell = hdr_cells[c_idx]
                    cell.width = COL_WIDTHS[c_idx]
                    set_cell_background(cell, "1B4D3E")  # Forest green
                    set_cell_margins(cell, top=120, bottom=120, left=120, right=120)
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    if c_idx == 4 or c_idx == 0:
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    r = p.add_run(h_text)
                    r.font.name = "Arial"
                    r.font.size = Pt(9)
                    r.font.bold = True
                    r.font.color.rgb = RGBColor(255, 255, 255)
                
                # Data Rows
                for r_idx, row_line in enumerate(rows_text):
                    # parse table cells: | **TC-FAR-001** | Feature | Instructions | Outcome | |
                    raw_cols = [c.strip() for c in row_line.split('|')[1:-1]]
                    data_cells = table.rows[r_idx + 1].cells
                    
                    # Alternating background
                    bg_color = "F9FAF9" if r_idx % 2 == 1 else "FFFFFF"
                    
                    for c_idx in range(5):
                        cell = data_cells[c_idx]
                        cell.width = COL_WIDTHS[c_idx]
                        set_cell_background(cell, bg_color)
                        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        
                        p = cell.paragraphs[0]
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(0)
                        
                        cell_content = raw_cols[c_idx] if c_idx < len(raw_cols) else ""
                        
                        if c_idx == 0:
                            # Test ID
                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            clean_id = cell_content.replace('**', '')
                            r = p.add_run(clean_id)
                            r.font.name = "Arial"
                            r.font.size = Pt(8.5)
                            r.font.bold = True
                            r.font.color.rgb = PRIMARY_COLOR
                        elif c_idx == 4:
                            # Pass / Fail blank column
                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r = p.add_run("")
                        else:
                            # Parse <br> and bold markers
                            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            lines = cell_content.split('<br>')
                            for l_idx, line in enumerate(lines):
                                if l_idx > 0:
                                    p = cell.add_paragraph()
                                    p.paragraph_format.space_before = Pt(1)
                                    p.paragraph_format.space_after = Pt(0)
                                
                                # parse bold **text**
                                parts = re.split(r'(\*\*.*?\*\*)', line)
                                for part in parts:
                                    if part.startswith('**') and part.endswith('**'):
                                        r = p.add_run(part[2:-2])
                                        r.font.bold = True
                                    else:
                                        r = p.add_run(part)
                                    r.font.name = "Arial"
                                    r.font.size = Pt(8.5)
                                    r.font.color.rgb = DARK_TEXT

    # Sign off section
    p_sign_h = doc.add_paragraph()
    p_sign_h.paragraph_format.space_before = Pt(18)
    p_sign_h.paragraph_format.space_after = Pt(6)
    p_sign_h.paragraph_format.keep_with_next = True
    r_sh = p_sign_h.add_run("Verification & Project Audit Sign-Off Table")
    r_sh.font.name = "Arial"
    r_sh.font.size = Pt(11)
    r_sh.font.bold = True
    r_sh.font.color.rgb = PRIMARY_COLOR

    sign_table = doc.add_table(rows=4, cols=4)
    sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(sign_table, color="D0D7DE")
    
    SIGN_WIDTHS = [Inches(1.8), Inches(2.2), Inches(2.0), Inches(1.3)]
    sign_headers = ["Project Role", "Name & Title", "Signature", "Date (DD/MM/YYYY)"]
    
    for c_idx, h_text in enumerate(sign_headers):
        cell = sign_table.rows[0].cells[c_idx]
        cell.width = SIGN_WIDTHS[c_idx]
        set_cell_background(cell, "2E7D32")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h_text)
        r.font.name = "Arial"
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        
    roles = [
        "Lead QA Automation Architect",
        "Lead Full-Stack Engineer",
        "System Auditor / Evaluator"
    ]
    for r_idx, role in enumerate(roles):
        row_cells = sign_table.rows[r_idx + 1].cells
        for c_idx in range(4):
            cell = row_cells[c_idx]
            cell.width = SIGN_WIDTHS[c_idx]
            set_cell_background(cell, "FFFFFF")
            set_cell_margins(cell, top=140, bottom=140, left=100, right=100)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if c_idx == 0:
                r = p.add_run(role)
                r.font.bold = True
            elif c_idx == 3:
                r = p.add_run("____ / ____ / 2026")
            else:
                r = p.add_run("")
            r.font.name = "Arial"
            r.font.size = Pt(8.5)
            r.font.color.rgb = DARK_TEXT

    output_path = "/Users/mevinimunaweera/Documents/EcoHarvest/ECOHARVEST_UAT_TEST_MANUAL.docx"
    doc.save(output_path)
    print(f"Successfully created: {output_path}")

if __name__ == "__main__":
    make_doc()
